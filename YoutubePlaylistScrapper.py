import requests
import os
from typing import List
import json
import csv
from docx import Document
from youtube_transcript_api import YouTubeTranscriptApi


class YoutubePlaylistScrapper:
    def __init__(self, playlist_link: str = None):
        playlist_id = json.load(open('C:\\Users\\jacob\\PycharmProjects\\fantano-knowledge-graph\\src\\main\\resources\\playlist_ids.json'))
        self.playlist_id = playlist_link
        self._scraped_video_ids = []
        self._scraped_video_info = []
        self.youtube_api_key = "NONE"

    def set_youtube_api_key(self, key: str) -> None:
        self.youtube_api_key = key

    def scrape(self):
        if self.youtube_api_key == "NONE":
            raise Exception("Youtube API key is not set. Please set the key in the YoutubePlaylistScrapper class.")
        self.scrape_video_ids()
        self.scrape_vidoes()
        self._add_transcript()


    def scrape_video_ids(self, next_page_token: str = None, total_results: int = -1, videos: List[str] = []) -> List[str]:

        address = f"https://www.googleapis.com/youtube/v3/playlistItems?playlistId={self.playlist_id}&key={self.youtube_api_key}&part=snippet&maxResults=50"

        if not next_page_token:
            vid_req = requests.get(address)

        else:
            vid_req = requests.get(address + f'&pageToken={next_page_token}')

        vids = vid_req.json()
        if "error" in vids.keys():
            raise Exception("Invalid API Key. Please check the API key and try again.")

        for x in vids['items']:
            videos.append(x['snippet']['resourceId']['videoId'])

        if total_results == -1:
            total_results = vids['pageInfo']['totalResults']
            print('total results set: ', total_results)

        print("total results: ", total_results)
        print("ids retrieved: ", len(videos), "\n")

        if "nextPageToken" not in vids.keys():
            print("processing")
            self._scraped_video_ids = videos
            return videos

        self.scrape_video_ids(next_page_token=vids['nextPageToken'], total_results=total_results, videos=videos)


    def scrape_vidoes(self):
        for video_id in self._scraped_video_ids:
            address = f"https://www.googleapis.com/youtube/v3/videos?id={video_id}&key={self.youtube_api_key}&part=snippet"
            vid_req = requests.get(address)
            vid = vid_req.json()
            self._scraped_video_info.append({"id": vid['items'][0]['id'], "title": vid['items'][0]['snippet']['title'],"url": f"https://www.youtube.com/watch?v={vid['items'][0]['id']}" ,"publish_date": vid['items'][0]['snippet']['publishedAt'][:10], "description": vid['items'][0]['snippet']['description']})


    def get_video_info(self) -> List[str]:
        return self._scraped_video_info


    def _create_transcript(self, video_id: str) -> str:
        try:
            raw_transcript = YouTubeTranscriptApi.get_transcript(video_id)
        except Exception:
            # Return an empty string if a transcript can't be fetched
            return "Captions are disabled for this video"

        transcript_formatted = ""
        for line in raw_transcript:
            # Convert the start time from seconds to hours, minutes, and seconds
            hours, remainder = divmod(line['start'], 3600)
            minutes, seconds = divmod(remainder, 60)

            # Format each line as: "HH:MM:SS - Text" or "MM:SS - Text" if hours is 00
            if hours == 0:
                transcript_formatted += f"{int(minutes):02d}:{int(seconds):02d} \n{line['text']}\n"
            else:
                transcript_formatted += f"{int(hours):02d}:{int(minutes):02d}:{int(seconds):02d} \n{line['text']}\n"

        return transcript_formatted

    def _add_transcript(self) -> None:
        for video in self._scraped_video_info:
            video['transcript'] = self._create_transcript(video['id'])

    def _write_to_csv(self) -> None:
        with open('transcripts/allTranscripts.csv', 'w', newline='', encoding='utf-8') as file:
            writer = csv.writer(file)
            writer.writerow(['id', 'title', 'publish_date', 'url','description','transcript'])
            for video in self._scraped_video_info:
                video['transcript'] = video['transcript'].replace('\n', ' ')
                writer.writerow([video['id'], video['title'], video['publish_date'], video['url'], video['description'], video['transcript']])
        print("CSV file created successfully.")


    def _write_to_docx(self) -> None:
        for video in self._scraped_video_info:
            document = Document()
            document.add_paragraph(video['url'])
            document.add_paragraph(video['publish_date'])
            document.add_paragraph(video['description'])
            document.add_paragraph('')  # Add a blank line
            document.add_paragraph(video['transcript'])
            invalid_characters = ['<', '>', ':', '"', '/', '\\', '|', '?', '*']
            title = video['title']
            for char in invalid_characters:
                title = title.replace(char, ' ')
            video['title'] = title
            document.save("transcripts/" + video['title'] + '.docx')
        print("Docx files created successfully.")